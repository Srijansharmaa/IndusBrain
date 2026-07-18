from pathlib import Path

from docx import Document as DocxDocument

from ai_engine.models.document import Document



from ai_engine.ingestion.base_reader import BaseReader

class DOCXReader(BaseReader):
    """
    Reads Word (.docx) files.
    """

    def __init__(self, file_path):

        self.file_path = Path(file_path)

    def read(self):

        doc = DocxDocument(self.file_path)

        text = []

        for paragraph in doc.paragraphs:

            if paragraph.text.strip():

                text.append(paragraph.text)

        return [

            Document(

                source=self.file_path.name,

                page=1,

                text="\n".join(text)

            )

        ]